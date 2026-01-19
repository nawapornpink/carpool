from __future__ import annotations

from io import BytesIO
from typing import Dict

from docx import Document
from docx.oxml import OxmlElement, ns
from docx.oxml.ns import qn


# =========================================================
# Token replace (รองรับ run แตก)
# =========================================================
def _replace_in_paragraph(paragraph, mapping: dict[str, str]) -> None:
    runs = paragraph.runs
    if not runs:
        return

    items = sorted(
        ((str(k), "" if v is None else str(v)) for k, v in mapping.items()),
        key=lambda kv: len(kv[0]),
        reverse=True,
    )

    def rebuild():
        texts = [r.text or "" for r in runs]
        full = "".join(texts)
        boundaries = []
        pos = 0
        for i, t in enumerate(texts):
            boundaries.append((pos, pos + len(t), i))
            pos += len(t)

        def run_at(char_pos: int) -> int:
            for s, e, idx in boundaries:
                if s <= char_pos < e:
                    return idx
            return boundaries[-1][2]

        return full, boundaries, run_at

    for token, value in items:
        if not token:
            continue

        while True:
            full, boundaries, run_at = rebuild()
            at = full.find(token)
            if at == -1:
                break

            token_start = at
            token_end = at + len(token)

            srun = run_at(token_start)
            erun = run_at(max(token_end - 1, token_start))

            srun_abs = boundaries[srun][0]
            erun_abs = boundaries[erun][0]

            s_text = runs[srun].text or ""
            e_text = runs[erun].text or ""

            left_keep = s_text[: max(token_start - srun_abs, 0)]
            right_keep = e_text[(token_end - erun_abs):]

            if srun == erun:
                runs[srun].text = left_keep + value + right_keep
            else:
                runs[srun].text = left_keep + value
                for i in range(srun + 1, erun):
                    runs[i].text = ""
                runs[erun].text = right_keep


def _replace_everywhere(doc: Document, mapping: dict[str, str]) -> None:
    for p in doc.paragraphs:
        _replace_in_paragraph(p, mapping)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph(p, mapping)

    for section in doc.sections:
        for p in section.header.paragraphs:
            _replace_in_paragraph(p, mapping)
        for p in section.footer.paragraphs:
            _replace_in_paragraph(p, mapping)


# =========================================================
# วาดเส้นขอบล่าง (ตัวจบเกม)
# =========================================================
def _set_bottom_border(cell, size: str = "8"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    bottom = OxmlElement("w:bottom")
    bottom.set(ns.qn("w:val"), "single")
    bottom.set(ns.qn("w:sz"), size)
    bottom.set(ns.qn("w:space"), "0")
    bottom.set(ns.qn("w:color"), "000000")

    tcBorders.append(bottom)


def _apply_mileage_borders(doc: Document, values: set[str]):
    """
    ถ้า cell.text ตรงกับค่าเลขไมล์ → ใส่เส้นขอบล่าง
    """
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text in values:
                    _set_bottom_border(cell)


# =========================================================
# Public API
# =========================================================
def build_car_docx(template_path: str, mapping: Dict[str, str]) -> BytesIO:
    """
    สร้างไฟล์ Word รายงานการใช้รถยนต์เช่า
    - replace token
    - วาดเส้นเลขไมล์จากโค้ด (ไม่พึ่ง Word)
    """
    doc = Document(template_path)

    # 1) replace token ทุกที่
    _replace_everywhere(doc, mapping)

    # 2) วาดเส้นเฉพาะช่องเลขไมล์
    mileage_values = set()
    if mapping.get("{{MILEAGE_START}}"):
        mileage_values.add(mapping["{{MILEAGE_START}}"])
    if mapping.get("{{MILEAGE_END}}"):
        mileage_values.add(mapping["{{MILEAGE_END}}"])

    _apply_mileage_borders(doc, mileage_values)

    # 3) return file
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio
